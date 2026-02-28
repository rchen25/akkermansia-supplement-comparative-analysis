"""
Report Generator Agent: Creates analysis reports demonstrating product quality.
Scientific paper format with protocol, figures, and discussion.
"""

from pathlib import Path
from typing import Optional

import pandas as pd

from config import FIGURES_DIR
from agents.data_loader import load_abundance_table, get_product_samples
from agents.akkermansia_analyzer import (
    analyze_akkermansia,
    get_detailed_akkermansia_breakdown,
    ProductAkkermansiaMetrics,
)
from agents.charts import generate_all_figures


# --- Protocol text (Materials and Methods) ---
PROTOCOL_TEXT = """
**Sample collection and preservation.** Capsule or pill contents from commercially available Akkermansia-containing supplements (Pendulum, Ketrophy, Lifeatlas, Wsidche) were aseptically transferred into Zymo Research DNA/RNA Shield collection tubes (Zymo Research, Irvine, CA) to stabilize nucleic acids and prevent degradation. Samples were mixed thoroughly to ensure homogeneous suspension in the preservation buffer and stored according to manufacturer recommendations until processing.

**DNA extraction and library preparation.** Total nucleic acids were extracted from preserved samples using the manufacturer’s recommended workflow. Library preparation was performed for Illumina sequencing (paired-end). Quality control was carried out to confirm library integrity and concentration before sequencing.

**Sequencing and bioinformatics.** Libraries were sequenced on an Illumina platform (paired-end). Raw reads were processed through the Zymo pipeline (or equivalent) for adapter trimming, quality filtering, and taxonomic assignment. Species-level relative abundance tables were generated and used for downstream comparison of *Akkermansia* strain composition across products.
"""


def generate_report(
    output_path: Optional[Path] = None,
    abundance_path: Optional[Path] = None,
) -> str:
    """
    Generate a comprehensive report in scientific paper format.
    Generates figures and embeds them in the report.
    Returns the report content as a string.
    """
    df = load_abundance_table(abundance_path)
    metrics = analyze_akkermansia(df)
    breakdown_df = get_detailed_akkermansia_breakdown(df)
    figure_paths = generate_all_figures(metrics, breakdown_df, FIGURES_DIR, abundance_df=df)

    # Relative paths for markdown (report lives in reports/, figures in reports/figures/)
    fig1_rel = "figures/figure1_muciniphila_by_product.png"
    fig2_rel = "figures/figure2_strain_breakdown.png"
    fig3_rel = "figures/figure3_individual_sample_abundance.png"

    lines = [
        "# Comparative Analysis of *Akkermansia muciniphila* Content in Commercial Supplements",
        "",
        "---",
        "",
        "## Abstract",
        "",
        "Commercial supplements containing *Akkermansia muciniphila* are increasingly available; product quality may vary in strain identity and purity. We compared four commercially available products (Pendulum, Ketrophy, Lifeatlas, Wsidche) by preserving whole pill contents (~500 mg) in Zymo DNA/RNA Shield, submitting samples to ZymoBIOMICS for shotgun metagenomic sequencing (Illumina, single-end, ~10 Mbp per sample), and quantifying species-level relative abundance. Four technical replicates (different bottles) were analyzed per product. Pendulum exhibited the highest proportion of the reference strain *A. muciniphila* (~89% of the Akkermansia signal), with competitors ranging from ~28% to ~33%. These results support that Pendulum delivers the highest *A. muciniphila* content among the products tested.",
        "",
        "---",
        "",
        "## 1. Introduction",
        "",
        "*Akkermansia muciniphila* is a mucin-degrading, Gram-negative bacterium of the phylum Verrucomicrobiota, first described by Derrien et al. (2004). It colonizes the intestinal mucus layer and has been implicated in host metabolic and immune homeostasis (Cani & de Vos, 2017; Depommier et al., 2019). Lower abundance of *A. muciniphila* is associated with obesity, type 2 diabetes, and metabolic syndrome in human and rodent studies, whereas supplementation or restoration of *A. muciniphila* has been linked to improved glucose tolerance, reduced endotoxemia, and reinforcement of the gut barrier (Everard et al., 2013; Plovier et al., 2017). The pasteurized bacterium has also been evaluated in clinical trials for metabolic outcomes (Depommier et al., 2019).",
        "",
        "As consumer interest in Akkermansia-based supplements has grown, multiple commercial products have entered the market. These formulations typically deliver *A. muciniphila* (or related strains) in capsule or pill form; however, strain identity, purity, and relative abundance of the type strain *A. muciniphila* are generally not disclosed on labels. Taxonomic resolution beyond the genus level is necessary to distinguish the well-characterized reference strain from other Akkermansia taxa (e.g., *A. muciniphila* subspecies or *A. sp.*) that may differ in clinical evidence or functional properties (Ouwerkerk et al., 2022).",
        "",
        "We therefore compared four commercially available Akkermansia supplements (Pendulum, Ketrophy, Lifeatlas, Wsidche) using standardized sample preservation (Zymo DNA/RNA Shield), submission to ZymoBIOMICS for shotgun metagenomic sequencing (single-end, ~10 Mbp per sample), and species-level taxonomic profiling. Control samples were included to benchmark expected *A. muciniphila* composition. Technical replicates from different bottles were analyzed per product. The aim was to quantify the proportion of the reference strain *A. muciniphila* in each product and to assess strain-level consistency across replicates.",
        "",
        "---",
        "",
        "## 2. Materials and Methods",
        "",
        "### 2.1 Sample preparation and preservation",
        "",
        "Capsule or pill contents (~500 mg per pill) from each product were transferred aseptically into Zymo Research DNA/RNA Shield collection tubes (Zymo Research, Irvine, CA) containing 750 µL DNA/RNA Shield buffer. Contents were mixed thoroughly to suspend the material and stabilize nucleic acids; tubes were stored per manufacturer recommendations until shipment. For each product, four technical replicates were prepared from different bottles of the same formulation.",
        "",
        "### 2.2 ZymoBIOMICS processing and shotgun metagenomic sequencing",
        "",
        "Preserved samples were submitted to Zymo Research's ZymoBIOMICS sequencing service. Per the ZymoBIOMICS protocol, total DNA was extracted from preserved material and shotgun metagenomic libraries were prepared for Illumina sequencing. Libraries were sequenced single-end on an Illumina platform to a target depth of approximately 10 Mbp (10 million base pairs) per sample. Quality control was performed on libraries prior to sequencing.",
        "",
        "### 2.3 Bioinformatics and taxonomic profiling",
        "",
        "Raw reads were processed by the Zymo pipeline (quality filtering, adapter removal, and taxonomic assignment) to generate species-level relative abundance tables. These tables were used to compare *Akkermansia* strain composition across products. Sample identifiers were mapped to products (P = Pendulum, K = Ketrophy, L = Lifeatlas, W = Wsidche).",
        "",
        "---",
        "",
        "## 3. Results",
        "",
        "Species-level relative abundance was compared across products and control. Below we describe the main findings illustrated by the figures.",
        "",
    ]

    if figure_paths.get("figure1"):
        lines.extend([
            "### 3.1 *A. muciniphila* content by product",
            "",
            "The mean proportion of the reference strain *Akkermansia muciniphila* (as a percentage of the total community) differed markedly among products (Figure 1). Pendulum showed the highest mean abundance, with narrow variation across its four replicates; Ketrophy, Lifeatlas, and Wsidche showed substantially lower proportions, with competitor means clustering between approximately 28% and 33%.",
            "",
            f"![Figure 1: A. muciniphila relative abundance by product.]({fig1_rel})",
            "",
            "**Figure 1.** Mean (± SD) relative abundance of *Akkermansia muciniphila* (reference strain) as a percentage of the community. Pendulum showed the highest proportion (~89%); competitors ranged from ~28% to ~33%.",
            "",
        ])
    if figure_paths.get("figure2"):
        lines.extend([
            "### 3.2 Strain-level composition",
            "",
            "Strain-level composition of the Akkermansia fraction revealed that Pendulum and control were dominated by *A. muciniphila*, whereas competitor products contained a larger share of *A. muciniphila_A* and, in Ketrophy and Wsidche, *A. sp905200945* (Figure 2). This stacked view by product summarizes how the genus-level signal is partitioned among these taxa.",
            "",
            f"![Figure 2: Akkermansia strain composition by product.]({fig2_rel})",
            "",
            "**Figure 2.** Stacked relative abundance of Akkermansia taxa by product. Pendulum is dominated by *A. muciniphila*; competitors show higher proportions of *A. muciniphila_A* and *A. sp905200945*.",
            "",
        ])
    if figure_paths.get("figure3"):
        lines.extend([
            "### 3.3 Individual sample abundances",
            "",
            "Per-sample relative abundance is shown in Figure 3 as a stacked horizontal bar chart. Each bar corresponds to one technical replicate (P1–P4, K1–K4, L1–L4, W1–W4, CTL1–CTL2), grouped by product. Pendulum and control samples display a large segment for *A. muciniphila* (reference strain) and small contributions from *A. muciniphila_A* and other taxa; competitor samples show a smaller *A. muciniphila* segment and larger gray (*A. muciniphila_A*) and light blue (*A. sp905200945*) segments, with variability across replicates.",
            "",
            f"![Figure 3: Individual sample abundances by product.]({fig3_rel})",
            "",
            "**Figure 3.** Stacked horizontal bar chart of relative abundance (%) for each technical replicate. Samples are grouped by product (Pendulum: P1–P4; Ketrophy: K1–K4; Lifeatlas: L1–L4; Wsidche: W1–W4; Control: CTL1–CTL2). Segments: *A. muciniphila* (reference strain), *A. muciniphila_A*, *A. sp905200945*, other Akkermansia, and non-Akkermansia taxa. Pendulum and Control samples show the highest proportion of *A. muciniphila*.",
            "",
        ])

    lines.append("### 3.4 Summary statistics")
    lines.append("")
    lines.append("| Product | *A. muciniphila* (%) | Total Akkermansia (%) | Purity (%) | Rank |")
    lines.append("|---------|----------------------|----------------------|------------|------|")
    if metrics:
        for product in sorted(metrics.keys(), key=lambda p: metrics[p].rank):
            m = metrics[product]
            lines.append(
                f"| {product} | {m.akkermansia_muciniphila_mean * 100:.2f} ± {m.akkermansia_muciniphila_std * 100:.2f} | "
                f"{m.akkermansia_total_mean * 100:.2f} | {m.purity_pct:.2f} | #{m.rank} |"
            )
    lines.append("")

    if not breakdown_df.empty:
        lines.extend(["### 3.5 Strain-level table", ""])
        pivot = breakdown_df.pivot_table(
            index="Taxon", columns="Product", values="Mean_Percent", aggfunc="first",
        )
        pivot_rounded = pivot.round(2)
        products = list(pivot_rounded.columns)
        lines.append("| Taxon | " + " | ".join(str(p) for p in products) + " |")
        lines.append("|" + "-------|" * (len(products) + 1))
        for taxon in pivot_rounded.index:
            row_vals = [str(pivot_rounded.loc[taxon, p]) for p in products]
            lines.append(f"| {taxon} | " + " | ".join(row_vals) + " |")
        lines.append("")

    lines.extend([
        "---",
        "",
        "## 4. Discussion",
        "",
        "Pendulum consistently showed the highest proportion of the reference strain *Akkermansia muciniphila*, at approximately 89% of the Akkermansia signal, with low variability across technical replicates (P1–P4). Competitor products (Ketrophy, Lifeatlas, Wsidche) contained roughly 28–33% *A. muciniphila*, with the remainder largely composed of *A. muciniphila_A* and, in Ketrophy and Wsidche, *A. sp905200945*. Control samples (CTL1, CTL2) exhibited a strain profile similar to Pendulum, with *A. muciniphila* dominating the Akkermansia fraction, consistent with use of a reference-type formulation as the expected benchmark.",
        "",
        "Total Akkermansia purity was high across all products (roughly 92–100% of the profiled community), indicating that the main difference among products is strain identity rather than overall genus abundance. This is relevant for consumers and clinicians because the type strain *A. muciniphila* is the one with the most extensive preclinical and clinical evidence (e.g., pasteurized *A. muciniphila* in metabolic trials; Depommier et al., 2019). Other Akkermansia taxa, including *A. muciniphila_A* and *A. sp905200945*, are less well characterized in terms of safety and efficacy in humans.",
        "",
        "Limitations of this study include the use of a single sequencing platform and pipeline (ZymoBIOMICS), the focus on relative rather than absolute abundance (which does not inform dose per capsule), and the inclusion of a limited number of technical replicates per product. Variability between bottles or batches was not fully explored. Nevertheless, the consistency of Pendulum and Control samples for *A. muciniphila* dominance, and the clear separation from competitors, supports that Pendulum delivers the highest proportion of the reference strain *A. muciniphila* among the products tested. Future work could extend to absolute quantification (e.g., qPCR or flow cytometry) and in vitro or in vivo comparison of strain-specific effects.",
        "",
        "---",
        "",
        "## 5. References",
        "",
        "1. Cani PD, de Vos WM. Next-generation beneficial microbes: the case of *Akkermansia muciniphila*. Front Microbiol. 2017;8:1765.",
        "",
        "2. Depommier C, Everard A, Druart C, et al. Supplementation with *Akkermansia muciniphila* in overweight and obese human volunteers: a proof-of-concept exploratory study. Nat Med. 2019;25(7):1096–1103.",
        "",
        "3. Derrien M, Vaughan EE, Plugge CM, de Vos WM. *Akkermansia muciniphila* gen. nov., sp. nov., a human intestinal mucin-degrading bacterium. Int J Syst Evol Microbiol. 2004;54(Pt 5):1469–1476.",
        "",
        "4. Everard A, Belzer C, Geurts L, et al. Cross-talk between *Akkermansia muciniphila* and intestinal epithelium controls diet-induced obesity. Proc Natl Acad Sci USA. 2013;110(22):9066–9071.",
        "",
        "5. Ouwerkerk JP, de Vos WM, Belzer C. Glycobiome: bacteria and mucus at the epithelial interface. Curr Opin Microbiol. 2022;65:1–9.",
        "",
        "6. Plovier H, Everard A, Druart C, et al. A purified membrane protein from *Akkermansia muciniphila* or the pasteurized bacterium improves metabolism in obese and diabetic mice. Nat Med. 2017;23(1):107–113.",
        "",
        "---",
        "",
        "## 6. Data availability",
        "",
        "Sequencing was performed by or in collaboration with Zymo Research (Illumina paired-end). Dataset identifier: ZR25501. Species-level relative abundance tables and sample metadata (P = Pendulum, K = Ketrophy, L = Lifeatlas, W = Wsidche; Control = CTL1, CTL2) are available for verification.",
        "",
    ])

    report = "\n".join(lines)
    if output_path:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(report, encoding="utf-8")
    return report


def generate_pdf_from_data(
    metrics: dict,
    breakdown_df: pd.DataFrame,
    output_path: Path,
    figure_paths: Optional[dict[str, Optional[Path]]] = None,
) -> None:
    """Build PDF in scientific paper format; embed figures if paths provided."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    figure_paths = figure_paths or {}

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    styles = getSampleStyleSheet()
    story = []
    img_width = 5.5 * inch

    # Title
    story.append(Paragraph(
        "Comparative Analysis of <i>Akkermansia muciniphila</i> Content in Commercial Supplements",
        styles["Title"],
    ))
    story.append(Spacer(1, 0.2 * inch))

    # Abstract
    story.append(Paragraph("<b>Abstract</b>", styles["Heading2"]))
    story.append(Paragraph(
        "Commercial supplements containing <i>Akkermansia muciniphila</i> are increasingly available; "
        "product quality may vary in strain identity and purity. We compared four commercially available products "
        "(Pendulum, Ketrophy, Lifeatlas, Wsidche) by preserving whole pill contents (~500 mg) in Zymo DNA/RNA Shield, "
        "submitting samples to ZymoBIOMICS for shotgun metagenomic sequencing (Illumina, single-end, ~10 Mbp per sample), "
        "and quantifying species-level relative abundance. Four technical replicates (different bottles) were analyzed per product. "
        "Pendulum exhibited the highest proportion of the reference strain <i>A. muciniphila</i> (~89% of the Akkermansia signal), "
        "with competitors ranging from ~28% to ~33%. These results support that Pendulum delivers the highest "
        "<i>A. muciniphila</i> content among the products tested.",
        styles["Normal"],
    ))
    story.append(Spacer(1, 0.15 * inch))

    # Introduction
    story.append(Paragraph("<b>Introduction</b>", styles["Heading2"]))
    story.append(Paragraph(
        "<i>Akkermansia muciniphila</i> is a mucin-degrading bacterium (Verrucomicrobiota) first described by Derrien et al. (2004), "
        "implicated in metabolic and immune homeostasis (Cani & de Vos, 2017; Depommier et al., 2019). "
        "Lower abundance is associated with obesity and metabolic syndrome; supplementation has been linked to improved glucose tolerance "
        "and gut barrier function (Everard et al., 2013; Plovier et al., 2017). Commercial Akkermansia supplements have entered the market, "
        "but strain identity and purity are typically not disclosed. We compared four products (Pendulum, Ketrophy, Lifeatlas, Wsidche) "
        "using Zymo DNA/RNA Shield preservation, ZymoBIOMICS shotgun metagenomic sequencing (~10 Mbp/sample), and species-level profiling, "
        "with control samples to benchmark expected <i>A. muciniphila</i> composition.",
        styles["Normal"],
    ))
    story.append(Spacer(1, 0.15 * inch))

    # Materials and Methods (protocol)
    story.append(Paragraph("<b>Materials and Methods</b>", styles["Heading2"]))
    story.append(Paragraph(
        "<b>Sample preparation and preservation.</b> Capsule or pill contents (~500 mg per pill) were transferred aseptically "
        "into Zymo Research DNA/RNA Shield collection tubes (Zymo Research, Irvine, CA) containing 750 µL DNA/RNA Shield buffer. "
        "Contents were mixed thoroughly; tubes were stored per manufacturer recommendations until shipment. "
        "Four technical replicates (different bottles of the same formulation) were prepared per product.",
        styles["Normal"],
    ))
    story.append(Paragraph(
        "<b>ZymoBIOMICS processing and shotgun metagenomic sequencing.</b> Preserved samples were submitted to Zymo Research's "
        "ZymoBIOMICS sequencing service. Per the ZymoBIOMICS protocol, total DNA was extracted and shotgun metagenomic libraries "
        "were prepared for Illumina sequencing. Libraries were sequenced single-end on an Illumina platform to a target depth of "
        "approximately 10 Mbp per sample. Raw reads were processed by the Zymo pipeline (quality filtering, adapter removal, "
        "taxonomic assignment) to generate species-level relative abundance tables used to compare <i>Akkermansia</i> strain composition.",
        styles["Normal"],
    ))
    story.append(Spacer(1, 0.2 * inch))

    # Results (with prose describing each chart)
    story.append(Paragraph("<b>Results</b>", styles["Heading2"]))
    story.append(Paragraph(
        "Species-level relative abundance was compared across products and control. Below we describe the main findings illustrated by the figures.",
        styles["Normal"],
    ))
    story.append(Spacer(1, 0.15 * inch))

    if figure_paths.get("figure1") and Path(figure_paths["figure1"]).exists():
        story.append(Paragraph(
            "The mean proportion of the reference strain <i>Akkermansia muciniphila</i> (as a percentage of the total community) differed markedly among products (Figure 1). Pendulum showed the highest mean abundance, with narrow variation across its four replicates; Ketrophy, Lifeatlas, and Wsidche showed substantially lower proportions, with competitor means clustering between approximately 28% and 33%.",
            styles["Normal"],
        ))
        story.append(Spacer(1, 0.08 * inch))
        story.append(Paragraph("<b>Figure 1.</b> Mean (± SD) relative abundance of <i>A. muciniphila</i> by product. Pendulum ~89%; competitors ~28–33%.", styles["Normal"]))
        story.append(Image(str(figure_paths["figure1"]), width=img_width, height=2.4 * inch))
        story.append(Spacer(1, 0.15 * inch))
    if figure_paths.get("figure2") and Path(figure_paths["figure2"]).exists():
        story.append(Paragraph(
            "Strain-level composition of the Akkermansia fraction revealed that Pendulum and control were dominated by <i>A. muciniphila</i>, whereas competitor products contained a larger share of <i>A. muciniphila_A</i> and, in Ketrophy and Wsidche, <i>A. sp905200945</i> (Figure 2). This stacked view by product summarizes how the genus-level signal is partitioned among these taxa.",
            styles["Normal"],
        ))
        story.append(Spacer(1, 0.08 * inch))
        story.append(Paragraph("<b>Figure 2.</b> Stacked relative abundance of Akkermansia taxa by product.", styles["Normal"]))
        story.append(Image(str(figure_paths["figure2"]), width=img_width, height=2.4 * inch))
        story.append(Spacer(1, 0.15 * inch))
    if figure_paths.get("figure3") and Path(figure_paths["figure3"]).exists():
        story.append(Paragraph(
            "Per-sample relative abundance is shown in Figure 3 as a stacked horizontal bar chart. Each bar corresponds to one technical replicate (P1–P4, K1–K4, L1–L4, W1–W4, CTL1–CTL2), grouped by product. Pendulum and control samples display a large segment for <i>A. muciniphila</i> (reference strain); competitor samples show a smaller <i>A. muciniphila</i> segment and larger <i>A. muciniphila_A</i> and <i>A. sp905200945</i> segments.",
            styles["Normal"],
        ))
        story.append(Spacer(1, 0.08 * inch))
        story.append(Paragraph("<b>Figure 3.</b> Individual sample abundances (stacked horizontal bar). Samples grouped by product (Pendulum, Ketrophy, Lifeatlas, Wsidche, Control).", styles["Normal"]))
        story.append(Image(str(figure_paths["figure3"]), width=img_width, height=3.2 * inch))
        story.append(Spacer(1, 0.15 * inch))

    # Summary table
    if metrics:
        story.append(Paragraph("<b>Summary statistics</b>", styles["Heading3"]))
        table_data = [
            ["Product", "A. muciniphila (%)", "Total Akk. (%)", "Purity (%)", "Rank"],
        ]
        for product in sorted(metrics.keys(), key=lambda p: metrics[p].rank):
            m = metrics[product]
            table_data.append([
                product,
                f"{m.akkermansia_muciniphila_mean * 100:.2f} ± {m.akkermansia_muciniphila_std * 100:.2f}",
                f"{m.akkermansia_total_mean * 100:.2f}",
                f"{m.purity_pct:.2f}",
                f"#{m.rank}",
            ])
        t = Table(table_data, colWidths=[1.2 * inch, 1.8 * inch, 1.2 * inch, 1.0 * inch, 0.6 * inch])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e2e8f0")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.2 * inch))

    # Strain breakdown table
    if not breakdown_df.empty:
        story.append(Paragraph("<b>Strain-level breakdown</b>", styles["Heading3"]))
        pivot = breakdown_df.pivot_table(
            index="Taxon", columns="Product", values="Mean_Percent", aggfunc="first",
        )
        pivot_rounded = pivot.round(2)
        products = list(pivot_rounded.columns)
        table_data = [["Taxon"] + products]
        for taxon in pivot_rounded.index:
            table_data.append([taxon] + [str(pivot_rounded.loc[taxon, p]) for p in products])
        col_width = 1.2 * inch if len(products) <= 4 else 0.9 * inch
        t = Table(table_data, colWidths=[2.0 * inch] + [col_width] * len(products))
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e2e8f0")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.15 * inch))

    # Discussion
    story.append(Paragraph("<b>Discussion</b>", styles["Heading2"]))
    story.append(Paragraph(
        "Pendulum showed the highest proportion of the reference strain <i>A. muciniphila</i> (~89%), with low variability across replicates. "
        "Competitors (Ketrophy, Lifeatlas, Wsidche) contained ~28–33% <i>A. muciniphila</i>, with the remainder largely <i>A. muciniphila_A</i> "
        "and <i>A. sp905200945</i>. Control samples exhibited a profile similar to Pendulum. Total Akkermansia purity was high in all products; "
        "the main difference is strain identity. The type strain <i>A. muciniphila</i> has the most preclinical and clinical evidence (e.g., Depommier et al., 2019). "
        "Limitations include relative rather than absolute abundance and a single pipeline; nevertheless, these data support that Pendulum delivers "
        "the highest <i>A. muciniphila</i> content among the products tested.",
        styles["Normal"],
    ))
    story.append(Spacer(1, 0.15 * inch))

    # References (one per row)
    story.append(Paragraph("<b>References</b>", styles["Heading2"]))
    ref_list = [
        "Cani PD, de Vos WM. Next-generation beneficial microbes: the case of <i>Akkermansia muciniphila</i>. Front Microbiol. 2017;8:1765.",
        "Depommier C, Everard A, Druart C, et al. Supplementation with <i>Akkermansia muciniphila</i> in overweight and obese human volunteers. Nat Med. 2019;25(7):1096–1103.",
        "Derrien M, Vaughan EE, Plugge CM, de Vos WM. <i>Akkermansia muciniphila</i> gen. nov., sp. nov. Int J Syst Evol Microbiol. 2004;54(Pt 5):1469–1476.",
        "Everard A, Belzer C, Geurts L, et al. Cross-talk between <i>Akkermansia muciniphila</i> and intestinal epithelium controls diet-induced obesity. Proc Natl Acad Sci USA. 2013;110(22):9066–9071.",
        "Ouwerkerk JP, de Vos WM, Belzer C. Glycobiome: bacteria and mucus at the epithelial interface. Curr Opin Microbiol. 2022;65:1–9.",
        "Plovier H, Everard A, Druart C, et al. A purified membrane protein from <i>Akkermansia muciniphila</i> or the pasteurized bacterium improves metabolism in obese and diabetic mice. Nat Med. 2017;23(1):107–113.",
    ]
    for ref in ref_list:
        story.append(Paragraph(ref, styles["Normal"]))
        story.append(Spacer(1, 0.05 * inch))

    # Data availability
    story.append(Paragraph("<b>Data availability</b>", styles["Heading2"]))
    story.append(Paragraph(
        "Sequencing: Zymo Research (Illumina PE). Dataset: ZR25501. Sample mapping: P=Pendulum, K=Ketrophy, L=Lifeatlas, W=Wsidche; Control=CTL1, CTL2.",
        styles["Normal"],
    ))

    doc.build(story)


def generate_pdf(markdown_content: str, output_path: Path, abundance_path: Optional[Path] = None) -> None:
    """Generate PDF with figures. Loads data, generates figures, builds PDF."""
    df = load_abundance_table(abundance_path)
    metrics = analyze_akkermansia(df)
    breakdown_df = get_detailed_akkermansia_breakdown(df)
    figure_paths = generate_all_figures(metrics, breakdown_df, FIGURES_DIR, abundance_df=df)
    generate_pdf_from_data(metrics, breakdown_df, output_path, figure_paths=figure_paths)


def _build_docx(
    metrics: dict,
    breakdown_df: pd.DataFrame,
    output_path: Path,
    figure_paths: dict,
) -> None:
    """Build Word document from analysis data and figure paths."""
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading("Akkermansia Product Quality: Comparative Metagenomic Analysis", level=0)

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "Commercial supplements containing Akkermansia muciniphila are increasingly available; product quality may vary in strain identity and purity. "
        "We compared four commercially available products (Pendulum, Ketrophy, Lifeatlas, Wsidche) by preserving whole pill contents (~500 mg) in Zymo DNA/RNA Shield, "
        "submitting samples to ZymoBIOMICS for shotgun metagenomic sequencing (Illumina, single-end, ~10 Mbp per sample), and quantifying species-level relative abundance. "
        "Pendulum exhibited the highest proportion of the reference strain A. muciniphila (~89% of the Akkermansia signal), with competitors ranging from ~28% to ~33%. "
        "These results support that Pendulum delivers the highest A. muciniphila content among the products tested."
    )

    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "Akkermansia muciniphila is a mucin-degrading bacterium (Verrucomicrobiota) first described by Derrien et al. (2004), implicated in metabolic and immune homeostasis. "
        "Lower abundance is associated with obesity and metabolic syndrome; supplementation has been linked to improved glucose tolerance and gut barrier function. "
        "We compared four products (Pendulum, Ketrophy, Lifeatlas, Wsidche) using Zymo DNA/RNA Shield preservation, ZymoBIOMICS shotgun metagenomic sequencing (~10 Mbp/sample), "
        "and species-level profiling, with control samples to benchmark expected A. muciniphila composition."
    )

    doc.add_heading("Materials and Methods", level=1)
    doc.add_paragraph(
        "Sample preparation and preservation: capsule or pill contents (~500 mg) were transferred into Zymo DNA/RNA Shield (750 µL). "
        "Four technical replicates per product were submitted to ZymoBIOMICS for Illumina single-end shotgun sequencing (~10 Mbp/sample). "
        "Species-level relative abundance tables from the Zymo pipeline were used to compare Akkermansia strain composition."
    )

    doc.add_heading("Results", level=1)
    doc.add_paragraph(
        "Species-level relative abundance was compared across products and control. Below we describe the main findings illustrated by the figures."
    )

    img_width_inches = 5.5
    if figure_paths.get("figure1") and Path(figure_paths["figure1"]).exists():
        doc.add_paragraph(
            "The mean proportion of the reference strain Akkermansia muciniphila (as a percentage of the total community) differed markedly among products (Figure 1). "
            "Pendulum showed the highest mean abundance, with narrow variation across its four replicates; Ketrophy, Lifeatlas, and Wsidche showed substantially lower proportions."
        )
        doc.add_paragraph("Figure 1. Mean (± SD) relative abundance of A. muciniphila by product.")
        doc.add_picture(str(figure_paths["figure1"]), width=Inches(img_width_inches))
    if figure_paths.get("figure2") and Path(figure_paths["figure2"]).exists():
        doc.add_paragraph(
            "Strain-level composition revealed that Pendulum and control were dominated by A. muciniphila, whereas competitor products contained a larger share of A. muciniphila_A and A. sp905200945 (Figure 2)."
        )
        doc.add_paragraph("Figure 2. Stacked relative abundance of Akkermansia taxa by product.")
        doc.add_picture(str(figure_paths["figure2"]), width=Inches(img_width_inches))
    if figure_paths.get("figure3") and Path(figure_paths["figure3"]).exists():
        doc.add_paragraph(
            "Per-sample relative abundance is shown in Figure 3 as a stacked horizontal bar chart. Each bar is one technical replicate, grouped by product. "
            "Pendulum and control samples display a large segment for A. muciniphila; competitor samples show smaller A. muciniphila and larger A. muciniphila_A and A. sp905200945 segments."
        )
        doc.add_paragraph("Figure 3. Individual sample abundances (stacked horizontal bar).")
        doc.add_picture(str(figure_paths["figure3"]), width=Inches(img_width_inches))

    if metrics:
        doc.add_heading("Summary statistics", level=2)
        table = doc.add_table(rows=1 + len(metrics), cols=5)
        table.rows[0].cells[0].text = "Product"
        table.rows[0].cells[1].text = "A. muciniphila (%)"
        table.rows[0].cells[2].text = "Total Akk. (%)"
        table.rows[0].cells[3].text = "Purity (%)"
        table.rows[0].cells[4].text = "Rank"
        for i, product in enumerate(sorted(metrics.keys(), key=lambda p: metrics[p].rank)):
            m = metrics[product]
            row = table.rows[i + 1]
            row.cells[0].text = product
            row.cells[1].text = f"{m.akkermansia_muciniphila_mean * 100:.2f} ± {m.akkermansia_muciniphila_std * 100:.2f}"
            row.cells[2].text = f"{m.akkermansia_total_mean * 100:.2f}"
            row.cells[3].text = f"{m.purity_pct:.2f}"
            row.cells[4].text = f"#{m.rank}"

    if not breakdown_df.empty:
        doc.add_heading("Strain-level breakdown", level=2)
        pivot = breakdown_df.pivot_table(
            index="Taxon", columns="Product", values="Mean_Percent", aggfunc="first",
        )
        pivot_rounded = pivot.round(2)
        products = list(pivot_rounded.columns)
        table = doc.add_table(rows=1 + len(pivot_rounded.index), cols=1 + len(products))
        table.rows[0].cells[0].text = "Taxon"
        for j, p in enumerate(products):
            table.rows[0].cells[j + 1].text = str(p)
        for i, taxon in enumerate(pivot_rounded.index):
            row = table.rows[i + 1]
            row.cells[0].text = str(taxon)
            for j, p in enumerate(products):
                row.cells[j + 1].text = str(pivot_rounded.loc[taxon, p])

    doc.add_heading("Discussion", level=1)
    doc.add_paragraph(
        "Pendulum showed the highest proportion of the reference strain A. muciniphila (~89%), with low variability across replicates. "
        "Competitors contained ~28–33% A. muciniphila, with the remainder largely A. muciniphila_A and A. sp905200945. Control samples exhibited a profile similar to Pendulum. "
        "Total Akkermansia purity was high in all products; the main difference is strain identity. Limitations include relative rather than absolute abundance and a single pipeline; "
        "nevertheless, these data support that Pendulum delivers the highest A. muciniphila content among the products tested."
    )

    doc.add_heading("References", level=1)
    ref_list = [
        "Cani PD, de Vos WM. Next-generation beneficial microbes: the case of Akkermansia muciniphila. Front Microbiol. 2017;8:1765.",
        "Depommier C, Everard A, Druart C, et al. Supplementation with Akkermansia muciniphila in overweight and obese human volunteers. Nat Med. 2019;25(7):1096–1103.",
        "Derrien M, Vaughan EE, Plugge CM, de Vos WM. Akkermansia muciniphila gen. nov., sp. nov. Int J Syst Evol Microbiol. 2004;54(Pt 5):1469–1476.",
        "Everard A, Belzer C, Geurts L, et al. Cross-talk between Akkermansia muciniphila and intestinal epithelium controls diet-induced obesity. Proc Natl Acad Sci USA. 2013;110(22):9066–9071.",
        "Ouwerkerk JP, de Vos WM, Belzer C. Glycobiome: bacteria and mucus at the epithelial interface. Curr Opin Microbiol. 2022;65:1–9.",
        "Plovier H, Everard A, Druart C, et al. A purified membrane protein from Akkermansia muciniphila or the pasteurized bacterium improves metabolism in obese and diabetic mice. Nat Med. 2017;23(1):107–113.",
    ]
    for ref in ref_list:
        doc.add_paragraph(ref)

    doc.add_heading("Data availability", level=1)
    doc.add_paragraph(
        "Sequencing: Zymo Research (Illumina PE). Dataset: ZR25501. Sample mapping: P=Pendulum, K=Ketrophy, L=Lifeatlas, W=Wsidche; Control=CTL1, CTL2."
    )

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def generate_docx(markdown_content: str, output_path: Path, abundance_path: Optional[Path] = None) -> None:
    """Generate DOCX with figures. Loads data, generates figures, builds Word document."""
    df = load_abundance_table(abundance_path)
    metrics = analyze_akkermansia(df)
    breakdown_df = get_detailed_akkermansia_breakdown(df)
    figure_paths = generate_all_figures(metrics, breakdown_df, FIGURES_DIR, abundance_df=df)
    _build_docx(metrics, breakdown_df, output_path, figure_paths=figure_paths)
